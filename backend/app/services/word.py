from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from sqlmodel import Session, select

from ..config import get_settings
from ..models import WordDocument
from ..schemas import WordDocumentRead
from .mistral_client import generate_summary
from .ocr import ensure_storage_dirs, get_converter

logger = logging.getLogger(__name__)


def documents_dir() -> Path:
    settings = get_settings()
    base = ensure_storage_dirs(settings.data_dir)["results"]
    directory = base / "word_documents"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def save_document(session: Session, document: WordDocument) -> WordDocument:
    document.updated_at = datetime.utcnow()
    session.add(document)
    session.commit()
    session.refresh(document)
    return document


def serialize_word_document(document: WordDocument, prefix: str) -> WordDocumentRead:
    return WordDocumentRead(
        id=document.id,
        title=document.title,
        source=document.source,
        original_filename=document.original_filename,
        file_name=document.file_name,
        mime_type=document.mime_type,
        summary=document.summary,
        created_at=document.created_at,
        download_url=f"{prefix}/word/documents/{document.id}/download",
    )


def generate_docx_from_text(title: str, content: str) -> Path:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    paragraphs = content.splitlines() or [content]
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    output_path = documents_dir() / f"{int(datetime.utcnow().timestamp())}_generated.docx"
    doc.save(output_path)
    return output_path


def create_word_document_from_text(session: Session, title: str, content: str) -> WordDocument:
    output_path = generate_docx_from_text(title, content)
    summary_input = (
        "Rezuma continutul urmatorului document Word in doua fraze in limba romana:\n" + content[:4000]
    )
    summary = generate_summary(summary_input)
    document = WordDocument(
        title=title or "Document fara titlu",
        source="generated",
        file_name=output_path.name,
        summary=summary,
    )
    return save_document(session, document)


def convert_pdf_to_word(session: Session, title: str, pdf_path: Path, original_filename: Optional[str] = None) -> WordDocument:
    converter = get_converter()
    result = converter.convert(str(pdf_path))
    markdown = result.document.export_to_markdown()
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    paragraphs = markdown.splitlines() or [markdown]
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    output_path = documents_dir() / f"{int(datetime.utcnow().timestamp())}_converted.docx"
    doc.save(output_path)
    summary_input = "Rezuma documentul convertit in doua fraze in limba romana:\n" + markdown[:4000]
    summary = generate_summary(summary_input)
    document = WordDocument(
        title=title or (original_filename or "Document convertit"),
        source="converted",
        original_filename=original_filename,
        file_name=output_path.name,
        summary=summary,
    )
    return save_document(session, document)


def get_document(session: Session, document_id: int) -> Optional[WordDocument]:
    return session.get(WordDocument, document_id)


def list_documents(session: Session) -> list[WordDocument]:
    statement = select(WordDocument).order_by(WordDocument.created_at.desc())
    return list(session.exec(statement))
